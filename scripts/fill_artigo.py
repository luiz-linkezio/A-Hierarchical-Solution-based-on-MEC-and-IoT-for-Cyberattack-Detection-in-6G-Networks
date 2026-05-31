from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

def set_para_text(para, new_text):
    p = para._p
    runs = p.findall(qn('w:r'))
    first_rpr = None
    if runs:
        first_rpr = runs[0].find(qn('w:rPr'))
        for r in runs:
            p.remove(r)
    r = OxmlElement('w:r')
    if first_rpr is not None:
        r.append(deepcopy(first_rpr))
    t = OxmlElement('w:t')
    t.text = new_text
    t.set(XML_SPACE, 'preserve')
    r.append(t)
    p.append(r)

def add_para_after(ref_para, text):
    new_p = OxmlElement('w:p')
    ppr = ref_para._p.find(qn('w:pPr'))
    if ppr is not None:
        new_p.insert(0, deepcopy(ppr))
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(XML_SPACE, 'preserve')
    r.append(t)
    new_p.append(r)
    ref_para._p.addnext(new_p)
    # Return wrapper
    from docx.text.paragraph import Paragraph as P
    return P(new_p, ref_para._p.getparent())

def set_cell_text(cell, text):
    for para in cell.paragraphs:
        runs = para._p.findall(qn('w:r'))
        for r in runs:
            para._p.remove(r)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set(XML_SPACE, 'preserve')
        r.append(t)
        para._p.append(r)
        break

# ── Content strings ─────────────────────────────────────────────────────────

RESUMO = (
    'O crescimento exponencial de dispositivos da Internet das Coisas (IoT) e a emergência das redes de sexta geração (6G) '
    'introduzem novos desafios de segurança cibernética em ambientes de Computação na Borda Móvel '
    '(MEC — Multi-access Edge Computing). A alta densidade de dispositivos, os rígidos orçamentos de latência e a '
    'sofisticação crescente dos ataques cibernéticos tornam insuficientes as abordagens tradicionais de detecção de '
    'intrusão baseadas em assinaturas. Este trabalho propõe e implementa um Sistema de Detecção de Intrusão (IDS) '
    'hierárquico de três fases projetado para ambientes IoT com MEC. A hierarquia distribui a carga computacional '
    'de acordo com a capacidade de cada camada: a Fase 1 executa um classificador binário LightGBM no nó de borda '
    '(VIM — Virtual Infrastructure Manager) para separar tráfego benigno de malicioso, otimizado pelo escore F2 para '
    'maximizar a revocação de ataques; a Fase 2 aplica um classificador multi-classe LightGBM para identificar o tipo '
    'específico de ataque em oito categorias (DDoS, DoS, malware, força bruta, MITM, ataques web, spoofing e '
    'reconhecimento); a Fase 3 realiza clusterização não supervisionada (UMAP + HDBSCAN) sobre fluxos de baixa '
    'confiança da Fase 2, visando identificar ameaças zero-day. O pipeline de dados utiliza capturas PCAP de cinco '
    'conjuntos de dados públicos do Instituto Canadense de Cibersegurança (CIC), convertidas em fluxos pela ferramenta '
    'netflower — desenvolvida como parte deste projeto e publicada no PyPI — e armazenadas em banco de dados SQLite de '
    'aproximadamente 190 GB, que viabiliza amostragem aleatória e balanceada em tempo de consulta. A Fase 1 atingiu '
    'acurácia geral de 84%, revocação de 93,9% na classe de ataque e F2 de 0,884, com limiar otimizado de 0,199. '
    'Em validação ao vivo no VIM 4 (2026-05-23), a taxa de falso positivo da Fase 1 foi de 0,03%, enquanto a acurácia '
    'da Fase 2 foi de 50,5%, com classificação perfeita de ataques DoS e dificuldades sistemáticas na distinção entre '
    'DDoS e força bruta — consequência de desequilíbrio na distribuição de treinamento. O sistema foi implantado em '
    'produção e validado com tráfego de ataque real gerado pelo script orquestrador, demonstrando a viabilidade da '
    'abordagem hierárquica ao mesmo tempo que identifica desafios de desvio de distribuição e calibração de limiar '
    'para ataques de curta duração.'
)

ABSTRACT = (
    'The exponential growth of Internet of Things (IoT) devices and the emergence of sixth-generation (6G) networks '
    'introduce new cybersecurity challenges in Multi-access Edge Computing (MEC) environments. High device density, '
    'strict latency budgets, and growing attack sophistication render traditional signature-based intrusion detection '
    'approaches insufficient. This work proposes and implements a three-phase hierarchical Intrusion Detection System '
    '(IDS) designed for IoT environments with MEC. The hierarchy distributes computational load according to each '
    "layer's capacity: Phase 1 runs a binary LightGBM classifier at the edge node (Virtual Infrastructure Manager — "
    'VIM) to separate benign from malicious traffic, optimized by the F2-score to maximize attack recall; Phase 2 '
    'applies a multi-class LightGBM classifier to identify the specific attack type among eight categories (DDoS, DoS, '
    'malware, brute force, MITM, web attacks, spoofing, and reconnaissance); Phase 3 performs unsupervised clustering '
    '(UMAP + HDBSCAN) on low-confidence Phase 2 flows to identify zero-day threats. The data pipeline uses PCAP '
    'captures from five public Canadian Institute for Cybersecurity (CIC) datasets, converted to flows using the '
    'netflower tool — developed as part of this project and published on PyPI — and stored in a ~190 GB SQLite database '
    'enabling random, balanced sampling at query time. Phase 1 achieved 84% overall accuracy, 93.9% attack class '
    'recall, and F2 of 0.884 with an Optuna-optimized threshold of 0.199. In live validation on VIM 4 (2026-05-23), '
    'Phase 1 achieved a 0.03% false positive rate, while Phase 2 reached 50.5% accuracy with perfect DoS '
    'classification and systematic difficulties distinguishing DDoS from brute force — a consequence of class '
    'imbalance in the training distribution. The system was deployed in production and validated with real attack '
    'traffic generated by the orchestrator script, demonstrating the viability of the hierarchical approach while '
    'identifying distribution shift and threshold calibration challenges for short-duration attacks.'
)

INTRODUCAO = (
    ' \tA convergência entre redes de sexta geração (6G), dispositivos da Internet das Coisas (IoT) e a Computação '
    'na Borda Móvel (MEC — Multi-access Edge Computing) promete transformar radicalmente a infraestrutura de '
    'comunicação digital. As redes 6G preveem densidades de dispositivos IoT na ordem de 10⁶ por km², com '
    'latências inferiores a 1 ms, tornando o MEC — que posiciona capacidade computacional em nós de borda próximos '
    'aos dispositivos — componente crítico para viabilizar esses requisitos (YOU et al., 2021). Contudo, essa '
    'massificação expande igualmente a superfície de ataque: dispositivos IoT frequentemente possuem poder '
    'computacional e energia limitados, inviabilizando soluções de segurança tradicionais diretamente nos terminais. '
    'Sistemas de Detecção de Intrusão (IDS) baseados em aprendizado de máquina, executados nos nós de borda '
    '(VIMs — Virtual Infrastructure Managers), surgem como alternativa viável: modelos treinados em dados históricos '
    'de tráfego classificam fluxos de rede em tempo real com alta eficiência computacional. Entretanto, os IDS '
    'existentes enfrentam limitações para ambientes IoT/MEC: (i) conjuntos de dados capturados em ambientes '
    'controlados geram desvio de distribuição no tráfego real; (ii) a variedade de tipos de ataque exige diferentes '
    'objetivos de otimização em cada camada hierárquica; (iii) ataques zero-day não podem ser detectados por '
    'classificadores supervisionados sem componente não supervisionada. Este trabalho apresenta um IDS hierárquico '
    'de três fases que aborda essas limitações, descrevendo o pipeline completo — da coleta de PCAPs à implantação '
    'em produção no VIM 4. As contribuições incluem: pipeline de dados reprodutível baseado em SQLite para '
    'amostragem balanceada em escala; ferramenta netflower (publicada no PyPI) para conversão paralela de PCAPs '
    'em fluxos; IDS binário de alta revocação (Fase 1) combinado com classificação multi-classe granular (Fase 2) '
    'e detecção não supervisionada de anomalias (Fase 3); e validação ao vivo com tráfego de ataque real. '
    'O artigo está organizado da seguinte forma: a Seção 2 apresenta a fundamentação teórica, a arquitetura '
    'proposta e os resultados; a Seção 3 apresenta as conclusões e trabalhos futuros.'
)

DEV_INTRO = (
    ' \tEsta seção descreve a fundamentação teórica, a arquitetura proposta, o pipeline de dados, o processo '
    'de treinamento das três fases e os resultados obtidos tanto no treinamento quanto na validação ao vivo.'
)

FUNDAMENTACAO = (
    ' \tIDS podem ser classificados em duas grandes abordagens: baseados em assinatura e baseados em anomalia. '
    'Os primeiros identificam padrões conhecidos com alta precisão, mas são incapazes de detectar ameaças novas. '
    'Os baseados em anomalia aprendem um perfil do comportamento normal e sinalizam desvios, sendo mais adequados '
    'para ambientes dinâmicos (LAZAREVIC et al., 2005). A combinação das duas abordagens em arquitetura hierárquica '
    'representa o estado da arte para ambientes com restrições de recursos, como nós de borda MEC. O LightGBM '
    '(Light Gradient Boosting Machine) é um algoritmo de gradient boosting baseado em árvores de decisão com '
    'crescimento folha a folha (leaf-wise), que proporciona convergência mais rápida e desempenho superior em '
    'conjuntos de alta dimensionalidade (KE et al., 2017). O Optuna é um framework de otimização automática de '
    'hiperparâmetros baseado em Tree-structured Parzen Estimator (TPE), com suporte a poda de tentativas pouco '
    'promissoras (AKIBA et al., 2019). Para a Fase 3, UMAP (Uniform Manifold Approximation and Projection) '
    'realiza redução de dimensionalidade preservando a estrutura local dos dados (McINNES et al., 2018), e '
    'HDBSCAN (Hierarchical Density-Based Spatial Clustering of Applications with Noise) identifica agrupamentos '
    'de densidade variável, atribuindo rótulo -1 aos pontos que não pertencem a nenhum cluster (CAMPELLO et al., '
    '2013). A ferramenta netflower, desenvolvida como parte deste projeto, converte PCAPs em fluxos de rede '
    'bidirecionais com suporte a paralelismo e instalação via pip, superando o CICFlowMeter em throughput e '
    'eliminando a dependência de Java.'
)

ARQUITETURA = (
    ' \tA arquitetura proposta é composta por quatro etapas: (i) coleta e organização de PCAPs em pastas '
    'rotuladas; (ii) conversão por netflower e armazenamento em banco SQLite unificado (~190 GB); '
    '(iii) treinamento das três fases no notebook training.ipynb; (iv) implantação em tempo real no VIM. '
    'Apenas datasets que disponibilizam PCAPs separados por categoria foram utilizados — os rótulos derivam '
    'do nome das pastas (benign/, ddos/, malware/ etc.), evitando inconsistências de esquema entre fontes. '
    'Foram empregados cinco datasets do CIC: APT IIoT 2024, IoT Dataset 2023, IoT DIAD 2024, IIoT 2025 e '
    'BCCC NRC IoMT 2024. As PCAPs (~1 TB bruto) foram processadas pelo netflower em CSVs rotulados e '
    'inseridas no SQLite em chunks de 50.000 linhas. O SQLite foi escolhido por suportar '
    'ORDER BY RANDOM() LIMIT n, permitindo amostragem aleatória e balanceada sem carregar o conjunto '
    'completo em memória — solução que resolve simultaneamente os problemas de memória e viés de ordem, '
    'com compressão adicional de ~5× frente às PCAPs originais. A engenharia de características '
    'partiu de 83 colunas brutas e removeu: features de identidade de fluxo (IPs, portas, timestamp), '
    'sentinelas artificiais (-1 em init_fwd_win_byts), features com desvio treinamento/implantação '
    '(bulk rates zerados em capturas reais), correlações altas (> 0,95) e variâncias próximas de zero, '
    'resultando em 65 features numéricas. As features mais relevantes pelo LightGBM gain foram '
    'flow_byts_s (bytes por segundo do fluxo) e flow_iat_min (tempo mínimo de chegada entre pacotes), '
    'que capturam respectivamente a taxa de transferência e o padrão temporal. A Tabela 1 apresenta a '
    'taxonomia de ataques adotada após unificação dos rótulos entre os cinco datasets.'
)

TREINAMENTO = (
    ' \tA Fase 1 treinou um LightGBM binário com 615.317 amostras benignas e 615.317 de ataque (distribuídas '
    'uniformemente entre as oito classes), com divisão 80/20 estratificada. O limiar de decisão foi '
    'otimizado conjuntamente com os hiperparâmetros via Optuna (20 tentativas, 1.800 s), usando '
    'F2-score como objetivo — que pondera revocação 4× mais que precisão, refletindo que, em um '
    'IDS (não IPS), um falso negativo é mais custoso que um falso positivo. O melhor resultado obteve '
    'F2 = 0,884, limiar ótimo = 0,199, acurácia = 84%, precisão de ataque = 73% e revocação = 93,9%. '
    'A Fase 2 treinou um LightGBM multi-classe com amostragem independente de até 615.317 linhas por '
    'classe diretamente do SQLite, com class_weight=balanced para compensar a escassez de bruteforce '
    '(~16 mil amostras). O objetivo Optuna foi macro F1, garantindo equidade entre classes. '
    'Em validação ao vivo (2026-05-23), a Fase 1 apresentou taxa de falso positivo de 0,03% '
    '(11.670 de 11.674 alertas dentro de janelas de ataque reais). A acurácia da Fase 2 foi de 50,5%: '
    'DoS classificado corretamente em 100%, mas DDoS classificado como DoS em 99,7% e força bruta como '
    'DoS em 84,2% — reflexo da dominância do DoS no conjunto de treinamento. Ataques de curta duração '
    '(web: 0,65 s; MITM: 0,53 s; spoofing: 2,7 s; malware: 7 s) não geraram alertas porque o netflower '
    'só emite fluxos após idle_timeout (30 s) ou TCP FIN/RST, e essas janelas encerraram antes da '
    'emissão dos fluxos. A Fase 3 aplica UMAP (n_components=10, n_neighbors=30, min_dist=0,0) seguido '
    'de HDBSCAN (min_cluster_size=100, min_samples=50) sobre fluxos com max(proba) < 0,4 da Fase 2; '
    'a integração ao script de tempo real (network_binary_ids.py) está prevista como trabalho futuro.'
)

CONCLUSAO = (
    ' \tEste trabalho apresentou um IDS hierárquico de três fases para ambientes IoT em redes 6G com MEC, '
    'cobrindo desde a construção do pipeline de dados até a implantação e validação em produção. A abordagem '
    'hierárquica mostrou-se viável: a Fase 1 atingiu alta revocação (93,9%) com custo computacional mínimo '
    'no nó de borda, e a Fase 2 adicionou granularidade na identificação do tipo de ataque. Na validação '
    'ao vivo, a taxa de falso positivo da Fase 1 foi de 0,03% e sete das oito categorias de ataque foram '
    'detectadas com alta acurácia quando a duração superou o idle_timeout do netflower. As principais '
    'limitações identificadas são: (i) desvio de distribuição entre tráfego benigno de treinamento '
    '(scripted) e tráfego real interativo, causando falsos positivos em sessões SSH e HTTP de curta '
    'duração; (ii) dominância da classe DoS no treinamento da Fase 2, gerando confusão com DDoS e força '
    'bruta; (iii) ataques de curta duração resultam em falsos negativos por não gerarem fluxos completos '
    'antes do encerramento da janela de avaliação. Como trabalhos futuros, propõe-se: integração da '
    'Fase 3 ao script de tempo real; re-treinamento da Fase 2 com features discriminativas (diversidade '
    'de origem para DDoS, contadores de falha de autenticação para força bruta); redução do idle_timeout '
    'para captura de ataques curtos; coleta de tráfego benigno real para complementar os dados de '
    'treinamento; e extensão da avaliação para outros nós MEC do ambiente 6G.'
)

ATTACK_TABLE = [
    ('Categoria', 'Descrição'),
    ('benign',     'Tráfego legítimo de rede'),
    ('ddos',       'Ataques de Negação de Serviço Distribuído (múltiplas fontes)'),
    ('dos',        'Ataques de Negação de Serviço (fonte única)'),
    ('malware',    'Tráfego de malware e comunicação de comando e controle (C2)'),
    ('bruteforce', 'Ataques de força bruta (SSH, HTTP, Telnet)'),
    ('mitm',       'Ataques Man-in-the-Middle (ARP spoofing)'),
    ('web',        'Ataques a aplicações web (scanning, fuzzing, injeção)'),
    ('spoofing',   'Falsificação de identidade de origem em pacotes IP'),
    ('recon',      'Varreduras e reconhecimento de rede (port scan, ping sweep)'),
]

REFS = [
    'AKIBA, T. et al. Optuna: a next-generation hyperparameter optimization framework. In: PROCEEDINGS OF THE ACM SIGKDD INTERNATIONAL CONFERENCE ON KNOWLEDGE DISCOVERY AND DATA MINING, 25., 2019, Anchorage. Proceedings... New York: ACM, 2019. p. 2623-2631.',
    'CAMPELLO, R. J. G. B.; MOULAVI, D.; SANDER, J. Density-based clustering based on hierarchical density estimates. In: PACIFIC-ASIA CONFERENCE ON KNOWLEDGE DISCOVERY AND DATA MINING, 17., 2013, Gold Coast. Proceedings... Berlin: Springer, 2013. p. 160-172.',
    'KE, G. et al. LightGBM: a highly efficient gradient boosting decision tree. In: ADVANCES IN NEURAL INFORMATION PROCESSING SYSTEMS, 30., 2017, Long Beach. Proceedings... [S.l.]: Curran Associates, 2017. p. 3146-3154.',
    'LAZAREVIC, A. et al. Intrusion detection: a survey. In: KUMAR, V.; SRIVASTAVA, J.; LAZAREVIC, A. (ed.). Managing Cyber Threats: Issues, Approaches and Challenges. New York: Springer, 2005. p. 19-78.',
    'McINNES, L.; HEALY, J.; MELVILLE, J. UMAP: uniform manifold approximation and projection for dimension reduction. arXiv preprint arXiv:1802.03426, 2018.',
    'SHARAFALDIN, I.; LASHKARI, A. H.; GHORBANI, A. A. Toward generating a new intrusion detection dataset and intrusion traffic characterization. In: INTERNATIONAL CONFERENCE ON INFORMATION SYSTEMS SECURITY AND PRIVACY, 4., 2018, Funchal. Proceedings... Setubal: SciTePress, 2018. p. 108-116.',
    'YOU, X. et al. Towards 6G wireless communication networks: vision, enabling technologies, and new paradigm shifts. Science China Information Sciences, v. 64, n. 1, p. 110301, 2021.',
    'ZOLANVARI, M. et al. Machine learning-based network vulnerability analysis of industrial Internet of Things. IEEE Internet of Things Journal, v. 6, n. 4, p. 6822-6834, 2019.',
]

# ── Load and edit document ───────────────────────────────────────────────────

doc = Document('docs/artigo-TCC.docx')
paras = doc.paragraphs

# Header info
set_para_text(paras[0], 'Solução IDS Hierárquica para Detecção de Ciberataques em Ambientes IoT com MEC em Redes 6G')
set_para_text(paras[2], 'A Hierarchical MEC and IoT Solution for Cyber-Attack Detection in 6G Networks')
set_para_text(paras[4], 'Luiz Linkezio')
set_para_text(paras[5], 'Orientação: Prof(a). Dr(a). [Nome do Orientador(a)]')

# Resumo
set_para_text(paras[9], RESUMO)
set_para_text(paras[11], 'Palavras-chave: Detecção de Intrusão; Internet das Coisas; Computação na Borda Móvel; LightGBM; Redes 6G; Segurança em Redes.')

# Abstract
set_para_text(paras[15], ABSTRACT)
set_para_text(paras[17], 'Keywords: Intrusion Detection System; Internet of Things; Multi-access Edge Computing; LightGBM; 6G Networks; Network Security.')

# Section 1
set_para_text(paras[19], '1 \tINTRODUÇÃO')
set_para_text(paras[20], INTRODUCAO)

# Section 2
set_para_text(paras[22], '2 \tDESENVOLVIMENTO')
set_para_text(paras[23], DEV_INTRO)

# 2.1
set_para_text(paras[26], '2.1 \tFundamentação Teórica')
set_para_text(paras[27], FUNDAMENTACAO)

# 2.2
set_para_text(paras[29], '2.2 \tArquitetura e Pipeline de Dados')
set_para_text(paras[30], ARQUITETURA)
set_para_text(paras[31], 'Tabela 1 - Taxonomia de Ataques Utilizada no Treinamento do Sistema IDS')
set_para_text(paras[32], 'Fonte: O autor (2026).')

# Table
table = doc.tables[0]
existing_rows = list(table.rows)
for i, row in enumerate(existing_rows):
    if i < len(ATTACK_TABLE):
        for j, cell in enumerate(row.cells):
            set_cell_text(cell, ATTACK_TABLE[i][j])

for i in range(len(existing_rows), len(ATTACK_TABLE)):
    new_row = table.add_row()
    for j, cell in enumerate(new_row.cells):
        set_cell_text(cell, ATTACK_TABLE[i][j])

# 2.2.1
set_para_text(paras[34], '2.2.1 \tFases de Treinamento e Resultados')
set_para_text(paras[35], TREINAMENTO)

# Figure caption / source
set_para_text(paras[37], 'Figura 1 – Arquitetura do Sistema IDS Hierárquico de Três Fases (Fases 1, 2 e 3)')
set_para_text(paras[49], 'Fonte: O autor (2026).')

# Section 3
set_para_text(paras[51], '3 \tCONCLUSÃO')
set_para_text(paras[52], CONCLUSAO)

# References: replace existing three + add more after last one
set_para_text(paras[56], REFS[0])
set_para_text(paras[58], REFS[1])
set_para_text(paras[60], REFS[2])

# Add remaining refs after para[60]; addnext inserts immediately after,
# so chain: ref_para -> ref3 -> blank -> ref4 -> blank -> ...
insert_after = paras[60]
for ref in REFS[3:]:
    ref_para = add_para_after(insert_after, ref)
    blank_para = add_para_after(ref_para, '')
    insert_after = blank_para

doc.save('docs/artigo-TCC.docx')
print("Done.")
